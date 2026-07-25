from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("outputs/unpublished_manuscript_and_academic_context_summary.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9E1EA", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_paragraph(paragraph, size=11, color="111827", bold=False, italic=False):
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold
        run.italic = italic


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.bold = True
    if level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(31, 78, 121)
    else:
        run.font.size = Pt(12.5)
        run.font.color.rgb = RGBColor(31, 78, 121)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.8)
    run.font.color.rgb = RGBColor(31, 41, 55)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.05)
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(31, 41, 55)


def add_callout(doc, title, body, fill="F3F7FB", accent="2F75B5"):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(0.10)
    table.columns[1].width = Inches(6.25)
    table.rows[0].cells[0].width = Inches(0.10)
    table.rows[0].cells[1].width = Inches(6.25)
    for cell in table.rows[0].cells:
        set_cell_border(cell, color="C7D3E0")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(table.rows[0].cells[0], accent)
    set_cell_shading(table.rows[0].cells[1], fill)
    p = table.rows[0].cells[1].paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(11.5)
    r.font.color.rgb = RGBColor(17, 24, 39)
    p2 = table.rows[0].cells[1].add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.name = "Calibri"
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor(55, 65, 81)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.75)
    table.columns[1].width = Inches(4.75)
    header = table.rows[0].cells
    header[0].text = "Area"
    header[1].text = "Summary"
    for cell in header:
        set_cell_shading(cell, "E8EEF5")
        set_cell_border(cell)
        set_cell_margins(cell)
        for p in cell.paragraphs:
            style_paragraph(p, size=10.2, bold=True, color="0B2545")
    for area, summary in rows:
        cells = table.add_row().cells
        cells[0].text = area
        cells[1].text = summary
        for cell in cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cell.paragraphs:
                style_paragraph(p, size=10, color="1F2937")
        set_cell_shading(cells[0], "F7F9FC")
    return table


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.8)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Unpublished Manuscript and Academic Context Summary")
    r.font.name = "Calibri"
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    sr = subtitle.add_run("Prepared as a concise explanatory note for graduate-program correspondence")
    sr.font.name = "Calibri"
    sr.font.size = Pt(10.5)
    sr.font.color.rgb = RGBColor(75, 85, 99)

    add_callout(
        doc,
        "Purpose of this document",
        "This note summarizes an unpublished manuscript and provides context for academic trajectory, GPA, and motivation for pursuing a second MSc. It is designed to communicate the research experience without sending the full manuscript draft.",
        fill="F8FAFC",
        accent="2F75B5",
    )

    add_heading(doc, "1. Current Academic Status", 1)
    add_body(
        doc,
        "I am currently completing my MSc degree and expect to graduate in [Month Year]. My graduate training began in a PhD-track research environment, where I was asked from the beginning to prioritize laboratory research and project development.",
    )
    add_body(
        doc,
        "The purpose of pursuing an additional MSc is not to restart my academic direction, but to strengthen formal training in the field that my thesis project ultimately clarified as my long-term interest: computational proteomics, chemoproteomics, LC-MS/MS data analysis, and quantitative bioinformatics.",
    )

    add_heading(doc, "2. Unpublished Manuscript Summary", 1)
    add_body(
        doc,
        "The manuscript was intended for publication but has not yet been submitted. Progress was delayed because of frequent personnel changes and a shift in the lab's research direction. For this reason, I am not sending the full draft, but the manuscript represents substantial research experience and a meaningful contribution to my graduate training.",
    )
    add_table(
        doc,
        [
            (
                "Working topic",
                "Development of an initial-rate data-selection algorithm for analyzing potent irreversible covalent inhibitors that violate the pseudo-first-order kinetic assumption.",
            ),
            (
                "Scientific gap",
                "Classical covalent-inhibitor kinetic analysis assumes inhibitor concentration is much higher than enzyme concentration. This assumption becomes difficult or impossible for very potent inhibitors that rapidly inactivate the enzyme at low nanomolar concentrations.",
            ),
            (
                "Analytical approach",
                "The manuscript describes a computational strategy to select appropriate early time-window data, evaluate meaningful observed rate constants, and improve kinetic parameter estimation when whole-curve fitting becomes misleading.",
            ),
            (
                "Example system",
                "The method was illustrated using a potent SARS-CoV-2 PLpro covalent inhibitor, ID-5-95, with kinetic characterization and mass-spectrometry-supported covalent labeling evidence.",
            ),
            (
                "Relevance to my training",
                "Although this project was not my final thesis direction, it reflects my exposure to quantitative analysis, covalent inhibitor characterization, assay data interpretation, and computational method development.",
            ),
        ],
    )

    add_heading(doc, "3. Context for MSc GPA", 1)
    add_body(
        doc,
        "My MSc GPA is approximately 3.0. I want to provide context because the transcript alone does not fully reflect how my graduate training was structured. I originally entered the program as a PhD student, and from the beginning I was expected to prioritize research productivity and lab work. As a result, my coursework performance was solid but not as strong as it could have been under a course-centered master's program.",
    )
    add_body(
        doc,
        "This explanation is not intended to excuse the GPA. Rather, it provides context: my graduate effort was heavily research-centered, and my strongest development occurred through research projects, computational analysis, manuscript preparation, and thesis work.",
    )

    add_heading(doc, "4. Reason for Leaving the PhD Track and Pursuing a Second MSc", 1)
    add_body(
        doc,
        "During my PhD training, the research direction of the group shifted from its initial mass-spectrometry-centered direction toward biological experiments and assay development. Through that experience, I realized that my interest in wet-lab biological assays was not strong enough to sustain the long and demanding path of a PhD in that direction.",
    )
    add_body(
        doc,
        "At the same time, my thesis project clarified the field I do want to pursue. My thesis focused on computational analysis of dose-resolved chemoproteomic data, including curve fitting, candidate prioritization, pathway analysis, and structural interpretation. This project is closely aligned with the professional and academic direction I hope to develop further.",
    )
    add_bullets(
        doc,
        [
            "My transition reflects a refined research direction rather than a loss of commitment to graduate study.",
            "The second MSc would allow me to build stronger formal training in computational proteomics, bioinformatics, and quantitative biological data analysis.",
            "My long-term goal is to work at the interface of LC-MS/MS-based proteomics, chemoproteomics, computational analysis, and data-driven target prioritization.",
        ],
    )

    add_heading(doc, "5. Suggested Short Email Explanation", 1)
    add_body(
        doc,
        "The following paragraph can be used in the email body if a shorter version is preferred:",
    )
    add_callout(
        doc,
        "Email-ready version",
        "Regarding publications, I have an unpublished manuscript draft from a project that was intended for submission, but the work was delayed because of changes in lab personnel and research direction. Since the manuscript has not been submitted, I am not sending the full draft, but I can provide a concise summary of the project and my research involvement. I also wanted to provide context for my MSc GPA. I originally entered the program as a PhD student and was asked to prioritize research from the beginning, so my coursework performance was solid but not as strong as it could have been in a course-centered program. During my PhD training, the lab's direction shifted from mass-spectrometry-based research toward biological experiments and assay development. Through that experience, I realized that my strongest and most sustainable interest is in computational proteomics, chemoproteomics, LC-MS/MS data analysis, and quantitative bioinformatics. My MSc thesis project is closely aligned with this direction, which is why I am now seeking a second MSc to strengthen my training in the field I hope to pursue long term.",
        fill="FFF7ED",
        accent="D96524",
    )

    add_heading(doc, "6. Tone to Preserve in Correspondence", 1)
    add_bullets(
        doc,
        [
            "Be factual and professional when mentioning lab personnel changes or research-direction changes.",
            "Avoid sounding negative toward the previous advisor or lab.",
            "Frame the second MSc as a focused next step based on clearer research alignment.",
            "Emphasize that the unpublished manuscript and thesis work demonstrate research maturity beyond what GPA alone shows.",
        ],
    )

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run("Prepared summary - not a manuscript submission")
    fr.font.name = "Calibri"
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(107, 114, 128)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
